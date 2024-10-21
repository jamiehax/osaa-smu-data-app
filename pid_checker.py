import streamlit as st
from langchain_openai import AzureChatOpenAI
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from pypdf import PdfReader
from docx import Document
import textract
import hmac


# check password
def check_password():
    """
    Returns `True` if the user had the correct password.
    """

    def password_entered():
        """
        Checks whether a password entered by the user is correct.
        """

        if hmac.compare_digest(st.session_state["pid_password"], st.secrets["pid_password"]):
            st.session_state["pid_password_correct"] = True
            del st.session_state["pid_password"]  # remove the password
        else:
            st.session_state["pid_password_correct"] = False

    # return True if the password is validated.
    if st.session_state.get("pid_password_correct", False):
        return True

    st.markdown("#### PID Page Password")

    st.text_input(
        "Password",
        placeholder="enter the PID page password...",
        on_change=password_entered,
        key="pid_password",
        label_visibility="collapsed"
    )

    if "pid_password_correct" in st.session_state:
        st.error("😕 Password incorrect")
    return False


if not check_password():
    st.stop()


# function to extract text from .pdf file
def extract_text_from_pdf(pdf_file):
    text = ""
    pdf_reader = PdfReader(pdf_file)
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

# function to extract text from .docx file
def extract_text_from_docx(docx_file):
    doc = Document(docx_file)
    # Extract paragraphs
    full_text = [para.text for para in doc.paragraphs]
    
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text.append(cell.text)

    return "\n".join(full_text)

# function to extract text from .doc file
def extract_text_from_doc(doc_file):
    text = textract.process(doc_file)
    return text.decode("utf-8")

# return the base criteria
def get_criteria():
    criteria = """
                ### 1. Coherence with Thematic Guidance and Priorities:

                **1.1 Alignment with Priorities**: Does the PID clearly align with the USG's thematic guidance and priorities?

                **1.2 Relevance to Clusters**: How well does the PID address the specific needs and objectives of the relevant thematic cluster(s)?

                **1.3 Consistency**: Are the proposed activities, outputs, outcomes, and objectives consistent with the overall strategic guidance?

                
                ### 2. Application of Value Chain:

                **2.1 Value Chain Integration**: Does the PID effectively integrate a value chain approach to achieve the desired outcomes and objectives?

                **2.2 Outcome Focus**: Are the outcomes and objectives clearly defined and achievable through the proposed value chain?

                **2.3 Activity-Outcome Linkage**: How well do the proposed activities link to the expected outputs and outcomes?

                
                ### 3. Intentionality and Target Audience:

                **3.1 Intentionality**: Is there a clear intention behind each proposed deliverable based on OSAA's definitions?

                **3.2 Target Audience**: Does the PID identify and address the needs of the target audience effectively?

                
                ### 4. Cross-Functional Collaboration:

                **4.1 Collaborative Efforts**: Does the PID outline clear strategies for cross-functional collaboration (policy, monitoring, coordination, advocacy, advisory)?

                **4.2 Cross-functional Coordination**: How well does the PID facilitate coordination between different teams and functions?

                
                ### 5. Delivery of Activities (deliverables table):

                **5.1 Role Clarity**: Are the roles and responsibilities of different teams and cluster leads clearly defined?

                **5.2 Activity Planning**: Are the planned activities well-structured, feasible, and intentional in targeting a specific audience according to the value chain?

                **5.3 Costing of Activities**: Are the costs associated with each activity included?

                **5.4 Timeline and Milestones**: Does the PID include a realistic timeline with clear milestones?

                
                ### 6. Monitoring and Evaluation:

                **6.1 Performance Indicators**: Are there clear performance indicators to measure success?

                
                ### 7. Gender Mainstreaming, Disability Inclusion, Multilingualism:

                **7.1 Inclusion**: Does the PID clearly outline how gender, disability, and multilingualism will be addressed?
            """
    
    return criteria

# return the thematic description for the given cluster
def get_thematic_description(cluster):

    thematic_descriptions = {
        1: 
        """

        ### Financing for Development:
        
        Access to financing, combat Illicit financial flows, enhancinginternational tax cooperation, engaging credit rating agencies, and reducing the cost ofremittances.

        Sustainability will not be reached by persisting on past formulas that conceive Africa as a continent in need of aid and debt relief,but only if African countries develop and manage their own sources of funding, and are considered as partners with an equal standing
        to other international stakeholders. Under this concept, financing for development is not any more an external factor and African economies are no longer relegated to be providers of raw materials, in a dependent relationship to their disadvantage. The contemporary African state exists with its alliance with foreign interests. Financing for development becomes a tool for the empowerment of Africa, and that addresses illicit financial flows to enable Africa to manage its development. Financing for development should trigger efforts for a greater coordination and cooperation around key issues and players, including the private sector, governments, international organizations and civil society. Among the different sources of financing that can increase African countries' domestic resource mobilization capacity, three have been selected for their potential contributions to the sustainability of funding provided a radical shift in paradigm: illicit financial flows, remittances and debt.

        Goals:
        a) Reverse illicit financial flows and promote international tax cooperation to increase domestic resource mobilization.
        b) Maximize the potential of remittances in financing for development.
        c) Address the impact of credit rating agencies on African countries' debt and their access to international capital markets.

        """,
        2: 
        """

        ### Sustainable Development to Promote Sustainable Peace:
        
        Inclusive and equal institutional practices by countering conflict economies and build cohesive diverse societies as a driving force for peace.

        National governments and local level stakeholders in Africa have the primary agency in sustaining peace on the continent in line with the African Union Silencing the Guns initiative. A comprehensive understanding of conflict drivers, including external and historic, is a prerequisite to sustainable peace. Only by addressing conflict drivers could societies meet the Sustainable Development Goals, in line with the UN's Sustaining Peace Agenda. To be effective, support to national actors should leverage the four pillars of the United Nations (peace and security, human rights, rule of law, development), in other words, effective support requires a nexus approach. Building resilient national capacities that can help counter structural inequalities, exclusion, and other drivers that undermine social cohesion should be at the front and center of the UN's support to Africa. If such drivers are neglected, they may over time lead to violent conflict and risk undermining the sustainable development goals.

        Goals:
        a) Achieve inclusive and equal institutional practices.
        b) Enable sustainable socioeconomic transformations to counter conflict economies.
        c) Build cohesive diverse societies as a driving force for sustainable development and peace.

        """,
        3: 
        """

        ### Democracy Resilience and Human Capital:
        
        Human capital at the center of policy making in Africa toward resilient societies.
        
        Building resilient societies in Africa requires to be cognizant of the challenges and opportunities brought about, among others, by climate change, globalization, urbanization, technological progress and demographic patterns, and one that actively innovates and implements strategies to address these systemic challenges and transform them into opportunities. In this regard, the core of a resilient country is its human capital. Societies with higher degrees of social protection and education systems are more prone to adapt to challenges, since a greater percentage of their population can play innovator roles. Consequently, promoting social inclusion through democracy and good governance and paying particular attention to women, youth and vulnerable populations is the best strategy to build resilient societies.

        Goals:
        a) Place human capital at the center of policy-making in Africa.
        b) Promote resilient societies through increased participation of women and the youth.
        c) Protect the most vulnerable: children, refugees, minorities, and persons with disabilities.

        """,
        4: 
        """

        ### Science Technology and Innovation:
        
        Closing the gap on digital literacy and digital divide; tackling intellectual property rights to achieve a leapfrog development.

        “A crisis is a terrible thing to waste,” as Paul Romer, Stanford economist, once said. The present unprecedented crisis provides impetus for challenging the business-as-usual scenario and implementing decisions that underpin more sustainable and resilient societies. Key results of data analysis show that on average over the period 2000-2018, the share of exports of ICT products is 0.9 percent of Africa's total trade. Africa is the continent with the lowest percentage in terms of import and export of ICT goods. The COVID-19 pandemic has demonstrated the stark digital divide and poor level of connectivity in Africa. Due to COVID-19 pandemic-driven school closures, Africa's youth lost access to education, opportunity, and livelihoods because technology became essential for study and work. African individuals, businesses, education institutions and governments need to be digitally enabled and connected. Beyond this, however, there is a challenge that is less frequently discussed - how Africa's low participation in the intellectual property arena has served as a disincentive to innovation on the continent, including in the biomedical field, digital technology, agriculture and in other areas. In 2018, Africa's share represented only 0.5 percent of the world's patent applications, compared with 66.8 percent in Asia, 19 percent in North America and 10.9 percent in Europe.

        Goals:
        a) Unleash the potential of Science, Technology, and Innovation for recovering better after COVID- 19.
        b) Leverage Science, Technology, and Innovation to leapfrog toward sustainability.
        c) Overcome the intellectual property barriers.

        """
        ,
        5: 
        """

        ### Industrialization, Demographic Dividend, and the African Continental Free Trade Area (AfCFTA):
        
        Harness the demographic dividend to stimulate industrialization through African Continental Free Trade Area.

        The establishment of the African Continental Free Trade Area (AfCFTA) provides opportunities for promoting value addition, industrialization, and economic diversification, as well as developing local and regional value chains. It will enable African countries to create much needed employment for the large number African youth, and generate resources for investment in human capital development, thereby harnessing the continent's demographic dividend. The COVID-19 pandemic has demonstrated the critical role of digital transformation in Africa. The COVID-19 responses provide an opportunity to accelerate application of practical digital solutions at scale, with impact on jobs and livelihoods.

        Goals:
        a) Promote access to Global and Regional Value Chains through industrialization.
        b) Implement the African continental Free Trade Area for realizing Africa's economic transformation.
        c) Harness the demographic dividend to stimulate industrialization.

        """,
        6: 
        """

        ### Energy and Climate Action:
        
        Africa's energy mix, climate change and green growth, energytransition, and environmental policies.

        Energy-sector bottlenecks and power shortages cost Africa 2-4 per cent of GDP annually, and thereby undermining sustainable economic growth, job creation and investment. Energy deficits reinforce poverty, especially for women and people in rural areas (World Economic Forum). Chronic under-supply of secure and affordable electricity is a barrier to growth, food security, job creation and poverty reduction. It limits Africa's prospects for industrialization and mechanization of productive sectors including agricultural value chains. Africa consumes only 3.1 percent of the world's energy output. Sub- regionally, southern Africa consumes 45 percent of the continent's energy and Northern Africa uses around 40%. According to the Africa Energy Outlook 2020, close to 770 million people, majority of them in sub-Saharan Africa, lack access to electricity, despite a recent decline in this figure resulting from increased grid connections and a rapid rise in the deployment of off-grid systems (IEA).

        Goals:
        a) Invest in energy production and sustainability.
        b) Balancing energy demands and climate change.
        c) Enhance investments in off-grid solutions to mitigate cost and increase efficiency in rural electrification.

        """
    }
        
    return thematic_descriptions.get(cluster, "No thematic description available for this cluster.")

# return the strategic guidance for the given cluster
def get_strategic_guidance(cluster):

    strategic_guidances = {
        1: 
        """

        The need to shift the paradigm of Finance for Development —a precondition to sustainable development— by tackling the finance paradox, making Domestic Resource Mobilization the game changer, and reconfiguring Africa's ownership of debt management.

        Africa faces a financing paradox, with a large development gap despite being a net lender to the world due to illicit financial flows surpassing foreign direct investment and official development assistance.

        The majority of Africa's development in 2020 was funded domestically, yet $500 - $600 billion remains untapped annually, while unfair borrowing costs and weak domestic resource mobilization perpetuate reliance on external finance.

        To overcome this, Africa must shift to an internally driven financing model, strengthening domestic resource mobilization and targeting transformative investment for greater financial control and sustainable development.

        """,
        2: 
        """

        Durable peace requires sustainable development - it is not the other way around. Today, Africa's instability stems from a huge deficit in delivering development in the past decades. Without understanding the root causes that have been preventing Africa from delivering first the MDGs and then the SDGs, causes of conflict will never be addressed. 

        Africa's peace and development challenges stem from a complex interplay of internal factors—such as ineffective political governance, economic disparities, poverty, inequality, and human rights issues—and external pressures like global competition for natural resources and the rise of international terrorist networks. These external forces often exacerbate local grievances, particularly in states with weak governance, where competition for resources can deepen frustrations among marginalized communities.

        The illicit extraction and trade of resources, facilitated by porous borders and informal markets, are especially harmful in countries lacking transparent and inclusive governance. This situation amplifies local discontent and undermines the legitimacy of the state, weakening its capacity to effectively address both internal and external challenges.

        To achieve durable peace and stability, African policymakers need to prioritize addressing internal governance failures, ensuring that resource exploitation benefits local communities, and fostering inclusive decision-making processes. By focusing on these internal factors and reducing external interference, Africa can strengthen state capacity, tackle the root causes of instability, and advance a development process that is both sustainable and led by African interests.

        """,
        3: 
        """

        Governance, Human capital, and Innovation: the enablers, the key success factors. 

        Africa's sustainable development relies on strong governance, human capital, and innovation, with achieving SDG 16 (focused on peace, inclusive societies, and accountable institutions) being essential. Strong institutions and equitable policies provide the foundation for addressing development challenges and turning them into opportunities for growth.

        Human capital is central to leveraging innovation for industrialization. African countries need to invest in science, technology, and innovation (STI) while scaling up STEM education and developing intellectual property ecosystems. This will foster local innovation, facilitate technology transfer, and build a skilled workforce, enabling industrialization and economic transformation, especially in agriculture and natural resource sectors.

        Policy frameworks must shift from reactive, charity-based approaches to proactive, forward-looking human capital strategies that prioritize education, health, and innovation. This transformation will unlock human capital as a driver of economic growth and build social resilience across the continent.

        Current service delivery gaps in education and health represent lost opportunities for Africa's progress. By leveraging technological innovation, African countries can address these deficiencies, creating a dynamic policy environment that places human capital at the center of development efforts, ensuring a resilient and prosperous future for the continent.        

        """,
        4: 
        """

        Governance, Human capital, and Innovation: the enablers, the key success factors. 

        Africa's sustainable development relies on strong governance, human capital, and innovation, with achieving SDG 16 (focused on peace, inclusive societies, and accountable institutions) being essential. Strong institutions and equitable policies provide the foundation for addressing development challenges and turning them into opportunities for growth.

        Human capital is central to leveraging innovation for industrialization. African countries need to invest in science, technology, and innovation (STI) while scaling up STEM education and developing intellectual property ecosystems. This will foster local innovation, facilitate technology transfer, and build a skilled workforce, enabling industrialization and economic transformation, especially in agriculture and natural resource sectors.

        Policy frameworks must shift from reactive, charity-based approaches to proactive, forward-looking human capital strategies that prioritize education, health, and innovation. This transformation will unlock human capital as a driver of economic growth and build social resilience across the continent.

        Current service delivery gaps in education and health represent lost opportunities for Africa's progress. By leveraging technological innovation, African countries can address these deficiencies, creating a dynamic policy environment that places human capital at the center of development efforts, ensuring a resilient and prosperous future for the continent.          

        """,
        5: 
        """

        Sustainable development requires high levels of resilience: the critical role of food systems - from unleashing small farmers to agri-processing driven by Africa's growing middle class - ISI the low-hanging fruit ). 

        Africa's industrialization is linked to its growing middle class and the shift towards import substitution (ISI), but past ISI efforts failed due to colonial economic structures focused on resource extraction rather than production.

        The expanding middle class offers a new opportunity to drive industrial growth by creating demand for value-added goods and services, necessitating a move away from commodity dependence.

        Continental import substitution can fast-track industrialization, job creation, and economic diversification, with the African Continental Free Trade Area (AfCFTA) playing a key role in promoting competitiveness and cross-border trade.

        Reliable energy access is crucial for industrialization, yet Africa faces an energy paradox—rich in resources but energy-deficient—making energy a critical component for powering industries and supporting sustainable development.

        To realize its industrial potential, Africa must leverage its middle class, adopt a "Made in Africa" approach, and build a strong energy infrastructure aligned with the demands of Industry 4.0.        

        """,
        6: 
        """

        Sustainable development requires full control of the nexus: energy planning - energy mix - energy access  

        Africa's energy paradox stems from its vast resources but limited access to reliable energy. Global climate discussions often ignore Africa's specific challenges in the energy transition, and despite climate pledges, rising CO2 emissions and a lack of accountability have restricted Africa's development and policy options.

        While Africa is committed to a green energy future, achieving it requires addressing its current energy deficits. Technological advancements in energy reliability, affordability, and distribution are critical to support both sustainable energy goals and Africa's economic and social needs.

        A balanced energy mix, strategic energy planning, and universal access are essential for Africa's future development, peace, and security. As energy demand grows with the continent's young population, these elements are crucial for a sustainable future.

        A one-size-fits-all energy transition is inadequate. Africa needs a "home-grown" approach that balances renewable energy with other sources to meet its development needs, avoiding an unrealistic reliance on renewables that could hinder its growth and energy independence.

        """,

    }
        
    return strategic_guidances.get(cluster, "No strategic guidance available for this cluster.")


# title and introduction
st.title("OSAA SMU's Check PID Tool")
st.markdown("Use the PID Checker to upload a PID and see if it aligns with the PID criteria.")

st.markdown("<hr>", unsafe_allow_html=True)
st.write("")


# initiatlize model
llm = AzureChatOpenAI(
    azure_deployment="osaagpt32k",
    api_key=st.secrets['azure'],
    azure_endpoint="https://openai-osaa-v2.openai.azure.com/",
    openai_api_version="2024-05-01-preview"
)


# upload PID
st.markdown("##### Upload a PID")
uploaded_file = st.file_uploader("Upload a PID", type=['pdf', 'doc', 'docx'], label_visibility="collapsed")
if uploaded_file is not None:

    # process file based on type
    if uploaded_file.type == "application/pdf":
        extracted_text = extract_text_from_pdf(uploaded_file)
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        extracted_text = extract_text_from_docx(uploaded_file)
    elif uploaded_file.type == "application/msword":
        extracted_text = extract_text_from_doc(uploaded_file)
    else:
        st.error("Unsupported file type. Please upload a DOC, DOCX, or PDF file type.")
        extracted_text = None
else:
    extracted_text = None


# select PID cluster
st.markdown("##### Select its Cluster")
selected_cluster = st.selectbox(
    "select a cluster...",
    (1, 2, 3, 4, 5, 6),
    index=None,
    placeholder="select a cluster...",
    label_visibility="collapsed"
    )


# get criteria based on cluster
criteria = get_criteria()
thematic_description = get_thematic_description(selected_cluster)
strategic_guidance = get_strategic_guidance(selected_cluster)

# chat prompt
prompt = ChatPromptTemplate.from_messages(
    [
        (
            "system",
            "You are an assistant. Your task is to determine whether the uploaded Project Initiation Document (PID) aligns with the criteria. The criteria is broken into 7 sections, each with its own sub-points. When evaluating a PID against the criteria, evaluate it section by section. Indicate where the PID aligns with the criteria and where it does not. Give recommendations on where the PID can improve to align with the criteria. Each PID belongs to a cluster, and you will also get a thematic description of the cluster and strategic guidance from the USG for the cluster that PID belongs to. Use the thematic description and strategic guidance to determine if the PID aligns with the criteria. The strategic guidance should be prioritized when evaluating the PID. If the PID aligns with the criteria based on the thematic description and strategic guidance for the cluster, clearly state that it does and provide reasons for each section and its sub-points. If the document does not align with the criteria based on the thematic description and strategic guidance for the cluster, clearly state that it does not and provide reasons for each section and its sub-points."
        ),
        (
            "human",
            "Project Initiation Document (PID): {document}"
        ),
        (
            "human",
            "criteria: {criteria}"
        ),
        (
            "human",
            "thematic description: {thematic_description}"
        ),
        (
            "human",
            "strategic guidance: {strategic_guidance}"
        )
    ]
)


# make chain
chain = (prompt | llm | StrOutputParser())


if st.button("Check PID", use_container_width=True, type="primary"):

    # get inputs
    inputs = {
        "document": extracted_text, 
        "criteria": criteria, 
        "thematic_description": thematic_description, 
        "strategic_guidance": strategic_guidance
    }

    # get response
    with st.spinner("checking PID..."):
        response = chain.invoke(inputs)

    st.markdown("#### Evaluation")
    st.write(response)

    # st.write(inputs)